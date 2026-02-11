import logging
import re
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any, Dict, List, Tuple
from zipfile import ZipFile

logger = logging.getLogger(__name__)

DOCX_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W_NS = DOCX_NS["w"]
P_TAG = f"{{{W_NS}}}p"
TBL_TAG = f"{{{W_NS}}}tbl"


def detect_images_in_docx(path: str) -> int:
    doc_path = Path(path)
    if doc_path.suffix.lower() != ".docx":
        return 0
    try:
        with ZipFile(doc_path, "r") as archive:
            return sum(1 for name in archive.namelist() if name.startswith("word/media/"))
    except Exception:
        return 0


def extract_docx_images(path: str, output_dir: str) -> List[str]:
    doc_path = Path(path)
    target_dir = Path(output_dir)
    target_dir.mkdir(parents=True, exist_ok=True)
    extracted: List[str] = []
    if doc_path.suffix.lower() != ".docx":
        return extracted

    try:
        with ZipFile(doc_path, "r") as archive:
            for name in archive.namelist():
                if not name.startswith("word/media/"):
                    continue
                raw = archive.read(name)
                out_name = Path(name).name
                out_path = target_dir / out_name
                out_path.write_bytes(raw)
                extracted.append(str(out_path))
    except Exception as exc:
        logger.warning("DOCX image extraction failed for %s: %s", doc_path, exc)
    return extracted


def _normalize_line(text: str) -> str:
    return " ".join(str(text or "").split()).strip()


def _xml_paragraph_text(paragraph_element: ET.Element) -> str:
    parts: List[str] = []
    for node in paragraph_element.iter():
        if node.tag == f"{{{W_NS}}}t":
            if node.text:
                parts.append(node.text)
            continue
        if node.tag == f"{{{W_NS}}}tab":
            parts.append(" ")
            continue
        if node.tag in {f"{{{W_NS}}}br", f"{{{W_NS}}}cr"}:
            parts.append("\n")
    text = "".join(parts)
    # Keep row-level line breaks but normalize each line.
    lines = [_normalize_line(line) for line in text.splitlines()]
    return "\n".join([line for line in lines if line]).strip()


def _xml_is_heading(paragraph_element: ET.Element) -> bool:
    p_style = paragraph_element.find("./w:pPr/w:pStyle", DOCX_NS)
    if p_style is None:
        return False
    style_val = str(p_style.attrib.get(f"{{{W_NS}}}val") or "").strip().lower()
    return style_val.startswith("heading")


def _iter_blocks(element: ET.Element):
    for child in list(element):
        if child.tag == P_TAG:
            yield ("p", child)
            continue
        if child.tag == TBL_TAG:
            yield ("tbl", child)
            continue
        # Recurse into wrappers (sdt/customXml/etc.) while preserving order.
        yield from _iter_blocks(child)


def _parse_docx_part_xml(
    root: ET.Element,
    *,
    include_headers: bool = False,
) -> Tuple[List[str], int, int, int, int, int]:
    lines: List[str] = []
    paragraph_count = 0
    table_count = 0
    table_cell_count = 0
    paragraphs_chars = 0
    table_chars = 0

    container = root.find("w:body", DOCX_NS) if not include_headers else root
    if container is None:
        return lines, paragraph_count, table_count, table_cell_count, paragraphs_chars, table_chars

    for kind, node in _iter_blocks(container):
        if kind == "p":
            paragraph_count += 1
            paragraph_text = _xml_paragraph_text(node)
            if not paragraph_text:
                continue
            paragraphs_chars += len(paragraph_text)
            if _xml_is_heading(node):
                lines.append(f"## {paragraph_text}")
            else:
                lines.append(paragraph_text)
            continue

        # kind == "tbl"
        table_count += 1
        lines.append("")
        rows = node.findall("./w:tr", DOCX_NS)
        if not rows:
            rows = node.findall(".//w:tr", DOCX_NS)
        for row in rows:
            row_cells: List[str] = []
            cells = row.findall("./w:tc", DOCX_NS)
            if not cells:
                cells = row.findall(".//w:tc", DOCX_NS)
            for cell in cells:
                table_cell_count += 1
                cell_lines: List[str] = []
                for paragraph in cell.findall(".//w:p", DOCX_NS):
                    paragraph_text = _xml_paragraph_text(paragraph)
                    if paragraph_text:
                        cell_lines.append(paragraph_text)
                cell_text = " / ".join(cell_lines).strip()
                table_chars += len(cell_text)
                row_cells.append(cell_text)
            lines.append("| " + " | ".join(row_cells) + " |")
        lines.append("")

    return lines, paragraph_count, table_count, table_cell_count, paragraphs_chars, table_chars


def _extract_docx_text_zip(path: Path) -> Tuple[str, Dict[str, Any]]:
    lines: List[str] = []
    paragraph_count = 0
    table_count = 0
    table_cell_count = 0
    paragraphs_chars = 0
    table_chars = 0

    with ZipFile(path, "r") as archive:
        if "word/document.xml" not in archive.namelist():
            return "", {
                "docx_paragraph_count": 0,
                "docx_table_count": 0,
                "docx_table_cell_count": 0,
                "extracted_from_tables_chars": 0,
                "extracted_from_paragraphs_chars": 0,
                "embedded_image_count": detect_images_in_docx(str(path)),
            }
        root = ET.fromstring(archive.read("word/document.xml"))
        main_lines, p_cnt, t_cnt, tc_cnt, p_chars, t_chars = _parse_docx_part_xml(root)
        lines.extend(main_lines)
        paragraph_count += p_cnt
        table_count += t_cnt
        table_cell_count += tc_cnt
        paragraphs_chars += p_chars
        table_chars += t_chars

        # Include auxiliary parts that often carry real recipe text in exported binders.
        aux_parts = [
            name
            for name in archive.namelist()
            if re.match(r"word/(header|footer)\d+\.xml$", name) or name in {"word/footnotes.xml", "word/endnotes.xml"}
        ]
        for part_name in aux_parts:
            try:
                part_root = ET.fromstring(archive.read(part_name))
            except Exception:
                continue
            aux_lines, p_cnt, t_cnt, tc_cnt, p_chars, t_chars = _parse_docx_part_xml(part_root, include_headers=True)
            lines.extend(aux_lines)
            paragraph_count += p_cnt
            table_count += t_cnt
            table_cell_count += tc_cnt
            paragraphs_chars += p_chars
            table_chars += t_chars

    cleaned_lines: List[str] = []
    for line in lines:
        normalized = _normalize_line(line)
        if not normalized:
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
            continue
        cleaned_lines.append(normalized)
    text = "\n".join(cleaned_lines).strip()
    return text, {
        "docx_paragraph_count": int(paragraph_count),
        "docx_table_count": int(table_count),
        "docx_table_cell_count": int(table_cell_count),
        "extracted_from_tables_chars": int(table_chars),
        "extracted_from_paragraphs_chars": int(paragraphs_chars),
        "embedded_image_count": int(detect_images_in_docx(str(path))),
    }


def _extract_docx_text_python_docx(path: Path) -> Tuple[str, Dict[str, Any]]:
    from docx import Document  # type: ignore

    document = Document(str(path))
    lines: List[str] = []
    paragraph_count = 0
    table_count = len(document.tables)
    table_cell_count = 0
    paragraphs_chars = 0
    table_chars = 0

    for paragraph in document.paragraphs:
        paragraph_count += 1
        text = (paragraph.text or "").strip()
        if not text:
            continue
        style_name = str(getattr(paragraph.style, "name", "") or "").lower()
        if style_name.startswith("heading"):
            lines.append(f"# {text}")
        else:
            lines.append(text)
        paragraphs_chars += len(text)

    for table in document.tables:
        lines.append("")
        for row in table.rows:
            row_cells: List[str] = []
            for cell in row.cells:
                table_cell_count += 1
                paragraphs = [(p.text or "").strip() for p in cell.paragraphs]
                cell_text = " / ".join(p for p in paragraphs if p).strip()
                table_chars += len(cell_text)
                row_cells.append(cell_text)
            lines.append("| " + " | ".join(row_cells) + " |")
        lines.append("")

    cleaned_lines: List[str] = []
    for line in lines:
        normalized = _normalize_line(line)
        if not normalized:
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
            continue
        cleaned_lines.append(normalized)
    text = "\n".join(cleaned_lines).strip()
    return text, {
        "docx_paragraph_count": int(paragraph_count),
        "docx_table_count": int(table_count),
        "docx_table_cell_count": int(table_cell_count),
        "extracted_from_tables_chars": int(table_chars),
        "extracted_from_paragraphs_chars": int(paragraphs_chars),
        "embedded_image_count": int(detect_images_in_docx(str(path))),
    }


def extract_text(path: str) -> Tuple[str, Dict[str, Any]]:
    file_path = Path(path)
    suffix = file_path.suffix.lower()

    if suffix == ".txt":
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return text, {
            "docx_paragraph_count": 0,
            "docx_table_count": 0,
            "docx_table_cell_count": 0,
            "extracted_from_tables_chars": 0,
            "extracted_from_paragraphs_chars": len(text),
            "embedded_image_count": 0,
        }

    if suffix != ".docx":
        return "", {
            "docx_paragraph_count": 0,
            "docx_table_count": 0,
            "docx_table_cell_count": 0,
            "extracted_from_tables_chars": 0,
            "extracted_from_paragraphs_chars": 0,
            "embedded_image_count": 0,
        }

    try:
        return _extract_docx_text_python_docx(file_path)
    except Exception as exc:
        logger.warning("python-docx extraction failed for %s; falling back to zip parser: %s", file_path, exc)
        return _extract_docx_text_zip(file_path)
